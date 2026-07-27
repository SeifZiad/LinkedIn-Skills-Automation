"""Generate the LinkedIn Skills Automation architecture Word guide."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, "Calibri", 16 if level == 1 else 13, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "Calibri", size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    for run in p.runs:
        set_run_font(run, "Calibri", 11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "Consolas", 10)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(10)
    # light gray shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F3F3")
    shd.set(qn("w:val"), "clear")
    pPr = p._p.get_or_add_pPr()
    pPr.append(shd)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, "Calibri", 11, bold=True)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = val
            for p in cells[c_i].paragraphs:
                for run in p.runs:
                    set_run_font(run, "Calibri", 11)
    doc.add_paragraph()


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LinkedIn Skills Automation Project")
    set_run_font(run, "Calibri", 22, bold=True, color=RGBColor(0x1A, 0x3A, 0x2F))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Complete Guide: Remote Debugger, SkillForge UI, Node.js Local Runner, and Localhost"
    )
    set_run_font(run, "Calibri", 12, color=RGBColor(0x55, 0x55, 0x55))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Project folder: LinkedIn_skills\n"
        "Audience: beginners who want a clear picture of how everything connects\n"
        "Document purpose: explain the idea, the why, and the exact startup flow"
    )
    set_run_font(run, "Calibri", 10, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # ---------- 1 ----------
    add_heading_styled(doc, "1. What problem this project solves", 1)
    add_para(
        doc,
        "This project automates adding skills to a LinkedIn profile. Instead of typing each skill "
        "manually and clicking education/experience checkboxes by hand, a Selenium test opens the "
        "LinkedIn “Add skill” form, types the skill name, selects related sources (for example Diploma, "
        "ISTQB, METI), saves, and continues with the next skill from a JSON data file.",
    )
    add_para(
        doc,
        "On top of that Java automation, there is a small browser UI called SkillForge. The UI is a "
        "friendly control panel: you choose which skills to add, optionally choose sources, press "
        "“Run automation”, and watch a live terminal-style log. The UI does not talk to LinkedIn by "
        "itself. It talks to a local helper program on your computer, and that helper starts the Java "
        "test which then controls Chrome.",
    )

    # ---------- 2 ----------
    add_heading_styled(doc, "2. The remote debugger method (attach to an existing Chrome)", 1)

    add_heading_styled(doc, "2.1 The simple idea", 2)
    add_para(
        doc,
        "Normally, when you write a Selenium test, Selenium launches a brand new Chrome window. That "
        "new window is clean: no cookies, no LinkedIn login, no browsing history from your normal "
        "profile. For many websites that is fine. For LinkedIn it is painful, because LinkedIn expects "
        "you to be logged in, and automating login can be blocked, slow, or unstable.",
    )
    add_para(
        doc,
        "The remote debugger method flips that approach. You open Chrome yourself first. You log into "
        "LinkedIn the normal human way. Then you tell Selenium: “Do not create a new browser. Connect "
        "to this Chrome that is already open and already logged in, and control it from the outside.”",
    )
    add_para(
        doc,
        "To allow that remote control, Chrome must be started with a special flag that opens a "
        "debugging “door” on a port number. In this project the door number is 9222.",
    )

    add_heading_styled(doc, "2.2 How Chrome is started for debugging", 2)
    add_para(doc, "On Windows PowerShell, a typical command looks like this:")
    add_code(
        doc,
        '& "C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe" '
        '--remote-debugging-port=9222 --user-data-dir="C:\\chrome-debug-profile"',
    )
    add_bullet(doc, "--remote-debugging-port=9222 means: listen for remote control commands on port 9222.")
    add_bullet(
        doc,
        "--user-data-dir=... means: use a separate Chrome profile folder for this debug session. "
        "That avoids conflicts with your everyday Chrome windows.",
    )
    add_para(
        doc,
        "After Chrome opens, you manually navigate to LinkedIn and sign in. Keep that window open "
        "while the test runs.",
    )

    add_heading_styled(doc, "2.3 How the Java code attaches", 2)
    add_para(
        doc,
        "In this project, WebDriverHandle configures ChromeOptions with an experimental option called "
        "debuggerAddress set to 127.0.0.1:9222. That tells Selenium to attach to the existing Chrome "
        "instead of launching a fresh one.",
    )
    add_para(doc, "Conceptually:")
    add_code(
        doc,
        "ChromeOptions options = new ChromeOptions();\n"
        "options.setExperimentalOption(\"debuggerAddress\", \"127.0.0.1:9222\");\n"
        "driver = new ChromeDriver(options);",
    )
    add_para(
        doc,
        "127.0.0.1 is another name for “this same computer”. So Selenium is saying: connect to the "
        "Chrome debug service on my own machine, port 9222.",
    )

    add_heading_styled(doc, "2.4 Why this method is used here", 2)
    add_bullet(doc, "LinkedIn requires an authenticated session.")
    add_bullet(doc, "A brand-new Selenium Chrome would force a login flow.")
    add_bullet(doc, "Manual login once, then automate the skill-adding clicks, is practical for demos and personal use.")
    add_bullet(doc, "It keeps your real session cookies and avoids repeatedly solving login challenges.")

    add_heading_styled(doc, "2.5 When you should use remote debugging", 2)
    add_bullet(doc, "When the website needs a real logged-in account and login automation is hard.")
    add_bullet(doc, "When you are demonstrating automation on your own machine.")
    add_bullet(doc, "When you want to inspect the page yourself in the same browser Selenium is controlling.")
    add_bullet(doc, "When cookies / 2FA / captcha make clean-browser login unreliable.")

    add_heading_styled(doc, "2.6 When you should not rely on it", 2)
    add_bullet(doc, "In company CI/CD pipelines that must start from a clean browser every time.")
    add_bullet(doc, "When tests must be fully unattended with no human login step.")
    add_bullet(doc, "When multiple machines need the same test without a pre-opened browser.")
    add_para(
        doc,
        "Remote debugging is powerful for local LinkedIn work, but it is a local convenience technique, "
        "not the only “correct” Selenium style for every environment.",
    )

    # ---------- 3 ----------
    add_heading_styled(doc, "3. How the SkillForge UI connects to the real LinkedIn profile", 1)

    add_heading_styled(doc, "3.1 Important truth first", 2)
    add_para(
        doc,
        "The HTML page cannot edit LinkedIn by itself. A webpage running in your browser is not allowed "
        "to freely control another website like LinkedIn and click around as you. That would be a huge "
        "security problem. So SkillForge is only a front panel. The real LinkedIn actions are done by "
        "your Java Selenium tests.",
    )

    add_heading_styled(doc, "3.2 The full chain of pieces", 2)
    add_para(doc, "End-to-end flow:")
    add_code(
        doc,
        "[1] SkillForge UI  (skills-dashboard.html in the browser)\n"
        "        |\n"
        "        |  HTTP request: POST /run with { skills: [...] }\n"
        "        v\n"
        "[2] Node.js local server  (ui/server.js on port 5050)\n"
        "        |\n"
        "        |  writes src/test/java/testdata/skills.JSON\n"
        "        |  starts: mvn test\n"
        "        v\n"
        "[3] Maven + TestNG + Selenium Java project\n"
        "        |\n"
        "        |  AddSkillsTest reads skills.JSON via DataProvider\n"
        "        |  profilepage types skills and clicks source checkboxes\n"
        "        v\n"
        "[4] Chrome already open with remote debugging on port 9222\n"
        "        |\n"
        "        v\n"
        "[5] Your real LinkedIn profile is updated",
    )

    add_heading_styled(doc, "3.3 What each piece does in detail", 2)

    add_para(doc, "SkillForge UI", bold=True)
    add_bullet(doc, "Section 1: queue the skill names you want to add.")
    add_bullet(doc, "Section 2: optionally choose sources (Diploma, ISTQB, METI, Ain Shams, and others). Sources clear after each Add so the next skill can use a fresh set.")
    add_bullet(doc, "Section 3: live terminal area that shows status lines while the run is happening.")
    add_bullet(doc, "Run automation button sends the queued skills to the local server.")

    add_para(doc, "Node.js local server", bold=True)
    add_bullet(doc, "Serves the HTML UI at http://localhost:5050/")
    add_bullet(doc, "Accepts POST /run")
    add_bullet(doc, "Validates the skills list")
    add_bullet(doc, "Overwrites skills.JSON in the Java test data folder")
    add_bullet(doc, "Runs mvn test in the project root")
    add_bullet(doc, "Streams Maven/TestNG/Selenium console output back to the UI terminal")

    add_para(doc, "Java automation layer", bold=True)
    add_bullet(doc, "HelperClass.ReadSkills loads skills.JSON into SkillData objects.")
    add_bullet(doc, "AddSkillsTest uses a TestNG DataProvider so each skill becomes its own test case.")
    add_bullet(doc, "profilepage opens the add-skill form URL, types the skill, presses Enter, selects sources, clicks Save.")
    add_bullet(doc, "Assertions check for the LinkedIn success message: “Your skill has been added”.")
    add_bullet(doc, "After the suite, the test can navigate to the skills details page.")

    add_para(doc, "Chrome + LinkedIn", bold=True)
    add_bullet(doc, "Because debuggerAddress is used, Selenium reuses the Chrome window where you are already logged in.")
    add_bullet(doc, "All visible clicking and typing happen in that real profile session.")

    add_heading_styled(doc, "3.4 What “connected to the real profile” means and does not mean", 2)
    add_para(
        doc,
        "It means the automation uses your authenticated LinkedIn browser session and performs the same "
        "UI actions a person would perform on the add-skill form. It does not mean the HTML page has a "
        "secret LinkedIn API key. There is no direct LinkedIn API call from the dashboard. The connection "
        "is: UI → local runner → Java test → Selenium → your open Chrome → LinkedIn website.",
    )

    # ---------- 4 ----------
    add_heading_styled(doc, "4. What is Node.js, and why do you open it before starting?", 1)

    add_heading_styled(doc, "4.1 Node.js in plain language", 2)
    add_para(
        doc,
        "Most people first meet JavaScript only inside web pages. Node.js is a runtime that lets "
        "JavaScript also run as a normal program on your computer, like Python or Java. With Node you "
        "can create small servers, read/write files, and start other programs.",
    )
    add_para(
        doc,
        "In this project, Node is used only for the local SkillForge helper (ui/server.js). Your "
        "automation logic for LinkedIn remains Java + Selenium. Node is the bridge between the pretty "
        "HTML UI and the Maven test command.",
    )

    add_heading_styled(doc, "4.2 Why you must start it before using the UI", 2)
    add_para(
        doc,
        "If you only open the HTML file from disk, the page can look nice, but the “Run automation” "
        "button has nowhere real to send the request (or it fails because nothing is listening). The "
        "Node server is the process that listens for that button click.",
    )
    add_para(doc, "When the server is running you typically see messages like:")
    add_code(
        doc,
        "SkillForge runner listening on http://localhost:5050\n"
        "Open UI: http://localhost:5050/\n"
        "Project root: D:\\ASU\\Testing\\LinkedIn_skills",
    )
    add_para(
        doc,
        "The UI also checks /health. If the server is up, the status shows something like "
        "“Runner connected · ready”. If Node is not running, it shows that the runner is offline.",
    )

    add_heading_styled(doc, "4.3 How to start the runner on this Windows machine", 2)
    add_para(
        doc,
        "If npm/node are not recognized in PowerShell because PATH is not refreshed, use the full path "
        "or add Node to PATH for the current session:",
    )
    add_code(
        doc,
        '$env:Path += ";C:\\Program Files\\nodejs"\n'
        "cd D:\\ASU\\Testing\\LinkedIn_skills\\ui\n"
        '& "C:\\Program Files\\nodejs\\node.exe" .\\server.js',
    )
    add_para(
        doc,
        "Keep that terminal window open while you use the dashboard. Closing it stops the local runner.",
    )
    add_para(
        doc,
        "Note about PowerShell execution policy: sometimes npm start fails because npm.ps1 is blocked. "
        "Calling node.exe server.js directly avoids that problem.",
    )

    # ---------- 5 ----------
    add_heading_styled(doc, "5. What is localhost?", 1)

    add_heading_styled(doc, "5.1 Localhost means “this computer”", 2)
    add_para(
        doc,
        "localhost is a special hostname that always points back to your own machine. When your browser "
        "opens http://localhost:5050, it is not going out to the public internet to find SkillForge. It "
        "is asking your own PC: “Is there a program listening on port 5050?” That program is the Node "
        "server.",
    )

    add_heading_styled(doc, "5.2 What the port number means", 2)
    add_para(
        doc,
        "A port is like a numbered door on your computer. Many programs can run at once, so each network "
        "service uses a different door number.",
    )
    add_table(
        doc,
        ["Address", "Who listens", "Purpose"],
        [
            [
                "http://localhost:5050",
                "Node.js SkillForge server",
                "Serve the UI and accept Run automation requests",
            ],
            [
                "127.0.0.1:9222",
                "Chrome started with remote debugging",
                "Allow Selenium to attach to the open browser",
            ],
        ],
    )
    add_para(
        doc,
        "127.0.0.1 and localhost are effectively the same idea: your own computer. Different ports are "
        "different services on that same computer.",
    )

    add_heading_styled(doc, "5.3 Is localhost public?", 2)
    add_para(
        doc,
        "By default this setup is local. Other people on the internet cannot open your SkillForge page "
        "just because you started it. It is for you, on your machine, for controlling your local test run.",
    )

    # ---------- 6 ----------
    add_heading_styled(doc, "6. Data flow for one skill run (very concrete)", 1)
    add_para(doc, "Example: you add skill “Java” with source “diploma”, then click Run automation.")
    add_bullet(doc, "UI stores in memory: { name: \"Java\", sources: [\"diploma\"] }")
    add_bullet(doc, "UI sends JSON to http://localhost:5050/run")
    add_bullet(doc, "server.js writes that into src/test/java/testdata/skills.JSON")
    add_bullet(doc, "server.js starts mvn test")
    add_bullet(doc, "TestNG DataProvider reads skills.JSON and creates one test invocation for “Java”")
    add_bullet(doc, "WebDriverHandle attaches to Chrome on 9222")
    add_bullet(doc, "profilepage opens the LinkedIn add-skill form URL")
    add_bullet(doc, "It types “Java”, presses Enter, scrolls to / clicks the diploma checkbox, clicks Save")
    add_bullet(doc, "It asserts the success toast appears")
    add_bullet(doc, "Console / UI terminal shows a success line for that skill")
    add_bullet(doc, "After all skills, the suite can open the skills details page")

    # ---------- 7 ----------
    add_heading_styled(doc, "7. JSON format used by both the UI and the Java tests", 1)
    add_para(
        doc,
        "The DataProvider update did not require a UI format change. Both sides share the same file shape:",
    )
    add_code(
        doc,
        '{\n'
        '  "skills": [\n'
        '    {\n'
        '      "name": "Software Testing",\n'
        '      "sources": ["diploma", "istqb", "meti"]\n'
        '    },\n'
        '    {\n'
        '      "name": "Java",\n'
        '      "sources": ["diploma"]\n'
        '    }\n'
        '  ]\n'
        '}',
    )
    add_para(
        doc,
        "Each object becomes one TestNG test via the DataProvider (name + sources passed as arguments). "
        "Sources are optional in the UI; if the list is empty, the page object skips checkbox selection.",
    )

    # ---------- 8 ----------
    add_heading_styled(doc, "8. Recommended startup order (checklist)", 1)
    add_para(doc, "Do these steps in order every time you want a real LinkedIn run:")
    add_para(doc, "Step 1 — Start debug Chrome and log in", bold=True)
    add_code(
        doc,
        '& "C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe" '
        '--remote-debugging-port=9222 --user-data-dir="C:\\chrome-debug-profile"',
    )
    add_para(doc, "Open LinkedIn and make sure you are logged into the correct account.")

    add_para(doc, "Step 2 — Start the Node runner", bold=True)
    add_code(
        doc,
        '$env:Path += ";C:\\Program Files\\nodejs"\n'
        "cd D:\\ASU\\Testing\\LinkedIn_skills\\ui\n"
        '& "C:\\Program Files\\nodejs\\node.exe" .\\server.js',
    )

    add_para(doc, "Step 3 — Open the UI", bold=True)
    add_code(doc, "http://localhost:5050/")
    add_para(doc, "Confirm the status says the runner is connected.")

    add_para(doc, "Step 4 — Queue skills and run", bold=True)
    add_bullet(doc, "Type a skill name.")
    add_bullet(doc, "Optionally select sources from the dropdown.")
    add_bullet(doc, "Click Add (sources clear automatically for the next skill).")
    add_bullet(doc, "Repeat for more skills.")
    add_bullet(doc, "Click Run automation.")
    add_bullet(doc, "Watch the live terminal until success or failure.")

    add_para(doc, "Alternative without the UI", bold=True)
    add_para(
        doc,
        "You can also edit skills.JSON by hand and run mvn test (or run AddSkillsTest from IntelliJ). "
        "Chrome must still be open on port 9222. The UI is optional convenience; the Java layer is the "
        "real automation engine.",
    )

    # ---------- 9 ----------
    add_heading_styled(doc, "9. Common beginner misunderstandings", 1)
    add_table(
        doc,
        ["Misunderstanding", "Reality"],
        [
            [
                "The HTML page edits LinkedIn directly",
                "No. HTML talks to Node; Node starts Java; Java/Selenium edits LinkedIn.",
            ],
            [
                "localhost is a website on the internet",
                "No. localhost means your own computer.",
            ],
            [
                "Node.js replaces Selenium",
                "No. Node only helps the UI and launches Maven. Selenium still drives Chrome.",
            ],
            [
                "Remote debugger opens LinkedIn for me",
                "No. You open Chrome and log in. Debugger only allows Selenium to attach.",
            ],
            [
                "One failed skill always means typing failed",
                "Often typing worked and the timeout is later, on source checkboxes or scrolling.",
            ],
        ],
    )

    # ---------- 10 ----------
    add_heading_styled(doc, "10. Glossary", 1)
    add_table(
        doc,
        ["Term", "Simple meaning"],
        [
            ["Selenium", "Library that can control a browser like a robot user"],
            ["WebDriver", "The bridge object your Java code uses to talk to Chrome"],
            ["debuggerAddress", "Setting that makes Selenium attach to an already-open Chrome"],
            ["Remote debugging port", "The door number Chrome opens for remote control (9222 here)"],
            ["Node.js", "Program that runs JavaScript on your computer"],
            ["Local runner / server.js", "Small Node app that serves UI and runs mvn test"],
            ["localhost", "Address that means “this PC”"],
            ["Port", "Numbered door for a network service on the PC"],
            ["skills.JSON", "Data file listing skills and sources for the tests"],
            ["DataProvider", "TestNG feature that turns each JSON skill into its own test"],
            ["Page Object (profilepage)", "Java class that knows LinkedIn page locators and actions"],
            ["Maven (mvn test)", "Build tool command that compiles and runs the tests"],
        ],
    )

    # ---------- 11 ----------
    add_heading_styled(doc, "11. Final picture in one paragraph", 1)
    add_para(
        doc,
        "You start Chrome with remote debugging so Selenium can reuse your logged-in LinkedIn session. "
        "You start a Node.js local server so the SkillForge HTML UI has a real backend on localhost:5050. "
        "When you press Run automation, the UI sends your skill list to that server; the server writes "
        "skills.JSON and runs your Maven/TestNG suite; each skill becomes one test through a DataProvider; "
        "Selenium attaches to Chrome on port 9222 and performs the LinkedIn add-skill actions on your "
        "real profile. That is the complete idea of this architecture.",
    )

    out = Path(__file__).resolve().parent / "LinkedIn_Skills_Automation_Architecture_Guide.docx"
    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
